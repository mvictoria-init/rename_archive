from pathlib import Path
import io
import html
from ebooklib import epub
from PyPDF2 import PdfReader
try:
    import fitz  # PyMuPDF
    _HAS_FITZ = True
except Exception:
    _HAS_FITZ = False
try:
    from PIL import Image
    import io
    _HAS_PIL = True
except Exception:
    _HAS_PIL = False
try:
    import docx
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False
try:
    from bs4 import BeautifulSoup
    _HAS_BS4 = True
except Exception:
    _HAS_BS4 = False


def pdf_to_epub(pdf_path, epub_path, title=None, authors=None):
    """Convert a PDF file to a simple EPUB.

    Args:
        pdf_path: path-like or str to source PDF
        epub_path: path-like or str to target EPUB
        title: optional title string
        authors: optional list of author strings

    Returns:
        (True, None) on success or (False, error_message) on failure.
    """
    try:
        pdf_path = Path(pdf_path)
        epub_path = Path(epub_path)

        reader = PdfReader(str(pdf_path))
        pages = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ''
            except Exception:
                txt = ''
            pages.append(txt)

        full = '\n\n'.join([p for p in pages if p])

        book = epub.EpubBook()
        if title:
            book.set_title(title)
        if authors:
            # authors may be list-like
            try:
                for a in authors:
                    book.add_author(a)
            except Exception:
                book.add_author(str(authors))

        # try to extract first-page image as cover (optional, requires PyMuPDF)
        reader = PdfReader(str(pdf_path))

        book = epub.EpubBook()
        if title:
            book.set_title(title)
        if authors:
            try:
                for a in authors:
                    book.add_author(a)
            except Exception:
                book.add_author(str(authors))

        # If PyMuPDF is available, attempt a richer extraction: headings, paragraphs and images.
        chapters = []
        # map PDF page index -> chapter index (used later to map TOC entries to chapter files)
        page_to_chap = {}
        cover_bytes = None
        if _HAS_FITZ:
            try:
                doc = fitz.open(str(pdf_path))
                # First, extract only the first page image as cover (JPEG for smaller size)
                try:
                    if doc.page_count > 0:
                        p0 = doc.load_page(0)
                        # render at a slightly higher resolution for decent results
                        pix = p0.get_pixmap(matrix=fitz.Matrix(2, 2))
                        try:
                            png_bytes = pix.tobytes('png')
                            if _HAS_PIL:
                                img = Image.open(io.BytesIO(png_bytes)).convert('RGB')
                                img.thumbnail((1600, 1600), Image.LANCZOS)
                                out = io.BytesIO()
                                img.save(out, format='JPEG', quality=75, optimize=True)
                                cover_bytes = out.getvalue()
                            else:
                                # fallback to the raw JPEG bytes from PyMuPDF if Pillow not available
                                cover_bytes = pix.tobytes('jpg')
                        except Exception:
                            # final fallback: try direct jpg bytes
                            try:
                                cover_bytes = pix.tobytes('jpg')
                            except Exception:
                                cover_bytes = None
                except Exception:
                    cover_bytes = None

                # Collect small snippets at top/bottom of each page to detect headers/footers
                page_blocks = []
                top_snippets = []
                bottom_snippets = []
                for pno in range(doc.page_count):
                    page = doc.load_page(pno)
                    blocks = page.get_text('dict').get('blocks', [])
                    parsed = []
                    for block in blocks:
                        if block.get('type') != 0:
                            continue
                        bbox = block.get('bbox', [])
                        y0 = bbox[1] if len(bbox) >= 4 else 0
                        max_sz = 0
                        texts = []
                        for line in block.get('lines', []):
                            for span in line.get('spans', []):
                                txt = span.get('text', '').strip()
                                if not txt:
                                    continue
                                sz = span.get('size', 0)
                                max_sz = max(max_sz, sz)
                                texts.append((sz, txt))
                        if not texts:
                            continue
                        txt_join = ' '.join(t for _, t in texts)
                        parsed.append((y0, max_sz, txt_join))
                    page_blocks.append((pno, parsed))
                    if parsed:
                        sorted_blocks = sorted(parsed, key=lambda x: x[0])
                        top_snippets.append(' | '.join(b[2] for b in sorted_blocks[:2]))
                        bottom_snippets.append(' | '.join(b[2] for b in sorted_blocks[-2:]))

                # determine frequent headers/footers (appear on >=50% pages)
                from collections import Counter
                def frequent(snips):
                    cnt = Counter(s for s in snips if s and len(s) > 3)
                    total = len(snips) or 1
                    return set(s for s, c in cnt.items() if c / total >= 0.5)

                common_tops = frequent(top_snippets)
                common_bottoms = frequent(bottom_snippets)

                # Build chapters: only text, no images; headings detected by font size
                for (pno, parsed) in page_blocks:
                    page_html = ''
                    for y0, max_sz, txt in parsed:
                        # filter likely header/footer
                        if txt in common_tops or txt in common_bottoms:
                            continue
                        if max_sz and max_sz >= 16:
                            page_html += f"<h2>{html.escape(txt)}</h2>"
                        else:
                            page_html += f"<p>{html.escape(txt).replace('\n', '<br/>')}</p>"
                    if page_html.strip():
                        chapters.append((pno, page_html))

                # attempt to extract table of contents (bookmarks)
                try:
                    raw_toc = doc.get_toc() or []
                except Exception:
                    raw_toc = []
                doc.close()
            except Exception:
                chapters = []

        # fallback: simple text extraction (one chapter)
        if not chapters:
            pages = []
            for page in reader.pages:
                try:
                    txt = page.extract_text() or ''
                except Exception:
                    txt = ''
                pages.append(txt)
            full = '\n\n'.join([p for p in pages if p])
            paragraphs = ''.join(f"<p>{html.escape(p).replace('\\n', '<br/>')}</p>" for p in full.split('\n\n') if p.strip())
            if title:
                chapters = [(0, f"<h1>{html.escape(title)}</h1>" + paragraphs)]
            else:
                chapters = [(0, paragraphs)]

        # build EPUB chapters/items
        spine = ['nav']
        chap_items = []
        # create EpubHtml items in the same order as chapters; remember mapping from PDF page -> chapter index
        for i, (orig_page, ch_html) in enumerate(chapters):
            fname = f'chap_{i+1}.xhtml'
            ctitle = title or f'Capítulo {i+1}'
            ch = epub.EpubHtml(title=ctitle, file_name=fname, lang='es')
            # content is expected to already contain basic tags (<h1>, <h2>, <p>) and no inline CSS
            ch.content = (ch_html or '')
            book.add_item(ch)
            chap_items.append(ch)
            spine.append(ch)
            if orig_page is not None:
                page_to_chap[orig_page] = i

        # cover
        if cover_bytes:
            try:
                book.set_cover('cover.jpg', cover_bytes)
            except Exception:
                pass

        # If we have a PDF TOC (bookmarks) from PyMuPDF, build a nested EPUB TOC mapping bookmark page -> chapter
        if _HAS_FITZ and 'raw_toc' in locals() and raw_toc:
            toc_entries = []
            stack = [(0, toc_entries)]
            # raw_toc entries are [level, title, page], page is 1-based
            for entry in raw_toc:
                try:
                    level, etitle, epage = entry
                except Exception:
                    continue
                target_page = int(epage) - 1
                # find nearest chapter for that page (search backward if exact page has no chapter)
                chap_idx = None
                for p in range(target_page, -2, -1):
                    if p in page_to_chap:
                        chap_idx = page_to_chap[p]
                        break
                if chap_idx is None:
                    # fallback to first chapter
                    chap_idx = 0 if chap_items else None
                if chap_idx is None:
                    continue
                ch = chap_items[chap_idx]
                link = epub.Link(ch.file_name, etitle or ch.title, ch.file_name)
                node_children = []
                node = (link, node_children)
                # adjust stack according to level
                while stack and level <= stack[-1][0]:
                    stack.pop()
                stack[-1][1].append(node)
                stack.append((level, node_children))
            book.toc = toc_entries
        else:
            book.toc = tuple((epub.Link(ch.file_name, ch.title, ch.file_name) for ch in chap_items))
            # Do not include 'nav' as the first spine item to avoid generating an extra TOC/title page
            book.spine = [s for s in spine if s != 'nav']
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        # write output
        epub_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            epub.write_epub(str(epub_path), book)
            return True, None
        except Exception as e:
            msg = str(e)
            if 'PyCryptodome' in msg or 'Crypto' in msg:
                return False, 'PyCryptodome is required for AES algorithm — please install with `pip install pycryptodome`'
            return False, msg
    except Exception as e:
        return False, str(e)


def _extract_text_from_docx(path):
    parts = []
    if not _HAS_DOCX:
        return parts
    try:
        doc = docx.Document(str(path))
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
    except Exception:
        return parts
    return parts


def _extract_cover_from_docx(path):
    if not _HAS_DOCX:
        return None
    try:
        doc = docx.Document(str(path))
        blobs = []
        for rel in getattr(doc.part, 'rels', {}).values():
            try:
                part = getattr(rel, 'target_part', None)
                if part is None:
                    continue
                blob = getattr(part, 'blob', None)
                if blob:
                    blobs.append(blob)
            except Exception:
                continue
        if not blobs:
            return None
        # choose largest image (likely cover)
        blob = max(blobs, key=len)
        if _HAS_PIL:
            try:
                img = Image.open(io.BytesIO(blob)).convert('RGB')
                img.thumbnail((1600, 1600), Image.LANCZOS)
                out = io.BytesIO()
                img.save(out, format='JPEG', quality=75, optimize=True)
                return out.getvalue()
            except Exception:
                return blob
        return blob
    except Exception:
        return None


def _extract_text_from_html(path):
    parts = []
    try:
        txt = Path(path).read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return parts
    if _HAS_BS4:
        try:
            soup = BeautifulSoup(txt, 'html.parser')
            # extract headings and paragraphs
            for tag in soup.find_all(['h1', 'h2', 'h3', 'p']):
                text = tag.get_text().strip()
                if text:
                    parts.append(text)
            return parts
        except Exception:
            pass
    # fallback simple tag stripping
    import re
    txt = re.sub(r'<[^>]+>', '\n', txt)
    for para in txt.split('\n'):
        p = para.strip()
        if p:
            parts.append(p)
    return parts


def _extract_text_from_txt(path):
    try:
        txt = Path(path).read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return []
    parts = [p.strip() for p in txt.split('\n\n') if p.strip()]
    return parts


def convert_to_epub(input_path, epub_path, title=None, authors=None):
    """Convert various document types to a lightweight EPUB.

    Supports: .pdf, .docx, .html/.htm, .txt. For PDF it prefers `pdf_to_epub`,
    but will fall back to text extraction if PDF parsing fails.
    """
    p = Path(input_path)
    suffix = p.suffix.lower()
    # If PDF, try our existing converter first
    if suffix == '.pdf':
        ok, err = pdf_to_epub(p, epub_path, title=title, authors=authors)
        if ok:
            return True, None
        # if PDF failed due to corruption, try fitz-based extraction or fallback to text
    # For other formats, extract text into chapters
    parts = []
    if suffix in ('.docx',) and _HAS_DOCX:
        # attempt to extract a cover image from docx
        cover_bytes = _extract_cover_from_docx(p)
        parts = _extract_text_from_docx(p)
    elif suffix in ('.html', '.htm'):
        parts = _extract_text_from_html(p)
    elif suffix in ('.txt',):
        parts = _extract_text_from_txt(p)
    else:
        # try best-effort: if it's a PDF but pdf_to_epub failed, try fitz extraction
        if _HAS_FITZ and suffix == '.pdf':
            try:
                doc = fitz.open(str(p))
                parts = []
                for page in doc:
                    t = page.get_text().strip()
                    if t:
                        parts.append(t)
                doc.close()
            except Exception:
                parts = []
        else:
            # unknown type: try reading as text
            parts = _extract_text_from_txt(p)

    if not parts:
        return False, 'No text extracted from input or unsupported format.'

    # build minimal chapters: split parts into chapters by detecting headings
    # Pre-filter parts: remove empty, duplicates, and parts equal to title/author
    clean_parts = []
    seen = None
    title_norm = (title or '').strip().lower()
    authors_text = ''
    if authors:
        try:
            authors_text = ' '.join(authors).strip().lower() if isinstance(authors, (list, tuple)) else str(authors).strip().lower()
        except Exception:
            authors_text = str(authors).strip().lower()
    for ptext in parts:
        if not ptext or not ptext.strip():
            continue
        s = ptext.strip()
        s_low = s.lower()
        if title_norm and s_low == title_norm:
            continue
        if authors_text and s_low == authors_text:
            continue
        if s == seen:
            continue
        clean_parts.append(s)
        seen = s

    chapters = []
    cur = []
    chap_idx = 0
    n = len(clean_parts)
    for i, block in enumerate(clean_parts):
        is_heading = False
        if len(block) < 120 and (block.isupper() or block.istitle() or block.endswith(':')):
            # check next block to ensure heading has content following
            next_block = clean_parts[i+1] if i+1 < n else ''
            if next_block and len(next_block) > 40:
                is_heading = True
        if is_heading:
            # flush current paragraph group
            if cur:
                ch_html = '<p>' + '</p><p>'.join(html.escape(x).replace('\n', '<br/>') for x in cur) + '</p>'
                chapters.append((chap_idx, ch_html))
                chap_idx += 1
                cur = []
            # add heading as its own chapter header (but only if not too repetitive)
            h = html.escape(block)
            if not chapters or (chapters and chapters[-1][1].find(h) == -1):
                chapters.append((chap_idx, f"<h2>{h}</h2>"))
                chap_idx += 1
        else:
            cur.append(block)
    if cur:
        ch_html = '<p>' + '</p><p>'.join(html.escape(x).replace('\n', '<br/>') for x in cur) + '</p>'
        chapters.append((chap_idx, ch_html))

    # build EPUB
    book = epub.EpubBook()
    if title:
        book.set_title(title)
    if authors:
        try:
            for a in authors:
                book.add_author(a)
        except Exception:
            book.add_author(str(authors))

    # attach cover if available (from PDF/docx extraction)
    try:
        if 'cover_bytes' in locals() and cover_bytes:
            try:
                book.set_cover('cover.jpg', cover_bytes)
            except Exception:
                pass
    except Exception:
        pass

    spine = ['nav']
    chap_items = []
    for i, (idx, ch_html) in enumerate(chapters):
        fname = f'chap_{i+1}.xhtml'
        ctitle = title or f'Capítulo {i+1}'
        ch = epub.EpubHtml(title=ctitle, file_name=fname, lang='es')
        ch.content = ch_html
        book.add_item(ch)
        chap_items.append(ch)
        spine.append(ch)

    book.toc = tuple((epub.Link(ch.file_name, ch.title, ch.file_name) for ch in chap_items))
    # Avoid placing the navigation page as the first spine entry (prevents an extra title/TOC page)
    book.spine = [s for s in spine if s != 'nav']
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub_path = Path(epub_path)
    epub_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        epub.write_epub(str(epub_path), book)
        return True, None
    except Exception as e:
        msg = str(e)
        if 'PyCryptodome' in msg or 'Crypto' in msg:
            return False, 'PyCryptodome is required for AES algorithm — please install with `pip install pycryptodome`'
        return False, msg
